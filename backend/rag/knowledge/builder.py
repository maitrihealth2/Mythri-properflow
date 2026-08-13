"""
Knowledge Base Builder (Multi-Format RAG)
Reads PDFs, DOCX, TXT, and JSON files from the docs/ directory.
Uses LangChain for text splitting and ChromaDB for vector storage.
"""

import os
import sys
import json
import shutil
import pathlib as _pl
from dotenv import load_dotenv

_BASE = _pl.Path(__file__).resolve().parent.parent.parent
load_dotenv(_BASE / ".env")
load_dotenv(_BASE / ".env.local", override=True)

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")

import chromadb
from chromadb.utils import embedding_functions

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    import docx
except ImportError:
    docx = None

try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
except ImportError:
    RecursiveCharacterTextSplitter = None

DOCS_DIR = os.path.join(os.path.dirname(__file__), "docs")
CHROMA_DIR = os.path.join(os.path.dirname(__file__), "chroma_db")
COLLECTION_NAME = "therapy_knowledge_v2"


def extract_text_from_pdf(filepath: str) -> str:
    if not fitz:
        print("PyMuPDF (fitz) is not installed. Skipping PDF.")
        return ""
    text = ""
    with fitz.open(filepath) as doc:
        for page in doc:
            text += page.get_text() + "\n"
    return text


def extract_text_from_docx(filepath: str) -> str:
    if not docx:
        print("python-docx is not installed. Skipping DOCX.")
        return ""
    doc = docx.Document(filepath)
    text = "\n".join([para.text for para in doc.paragraphs])
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_data:
                text += "\n" + " | ".join(row_data)
    return text


def chunk_text(text: str) -> list[str]:
    if RecursiveCharacterTextSplitter:
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=100,
            separators=["\n\n", "\n", ".", " ", ""]
        )
        return splitter.split_text(text)
    else:
        # Fallback simple chunker
        chunks = []
        chunk_size = 500
        overlap = 100
        start = 0
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            start += chunk_size - overlap
        return chunks


def load_documents() -> list[dict]:
    documents = []
    
    if not os.path.exists(DOCS_DIR):
        print(f"[RAG] Directory not found: {DOCS_DIR}")
        return documents
        
    for root, _, files in os.walk(DOCS_DIR):
        for filename in files:
            filepath = os.path.join(root, filename)
            rel_path = os.path.relpath(filepath, DOCS_DIR)
            source_name = os.path.splitext(filename)[0]
            
            if filename.endswith(".json"):
                with open(filepath, "r", encoding="utf-8") as f:
                    try:
                        data = json.load(f)
                        if isinstance(data, list):
                            for idx, chunk in enumerate(data):
                                documents.append({
                                    "id": f"{source_name}_{idx}",
                                    "text": chunk.get("text", ""),
                                    "source": source_name,
                                    "concept": chunk.get("concept", "general"),
                                    "technique": chunk.get("technique", "none")
                                })
                            print(f"  📄 {rel_path} → {len(data)} structured chunks")
                    except json.JSONDecodeError:
                        print(f"  ❌ Error parsing {rel_path}")

            elif filename.endswith(".txt") or filename.endswith(".md"):
                with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                    content = f.read()
                    chunks = chunk_text(content)
                    for idx, chunk in enumerate(chunks):
                        documents.append({
                            "id": f"{source_name}_txt_{idx}",
                            "text": chunk,
                            "source": source_name,
                            "concept": source_name,
                            "technique": "general"
                        })
                    print(f"  📄 {rel_path} → {len(chunks)} text chunks")

            elif filename.endswith(".pdf"):
                content = extract_text_from_pdf(filepath)
                if content.strip():
                    chunks = chunk_text(content)
                    for idx, chunk in enumerate(chunks):
                        documents.append({
                            "id": f"{source_name}_pdf_{idx}",
                            "text": chunk,
                            "source": source_name,
                            "concept": "clinical_knowledge",
                            "technique": "general"
                        })
                    print(f"  📄 {rel_path} → {len(chunks)} PDF chunks")

            elif filename.endswith(".docx"):
                content = extract_text_from_docx(filepath)
                if content.strip():
                    chunks = chunk_text(content)
                    for idx, chunk in enumerate(chunks):
                        documents.append({
                            "id": f"{source_name}_docx_{idx}",
                            "text": chunk,
                            "source": source_name,
                            "concept": "clinical_knowledge",
                            "technique": "general"
                        })
                    print(f"  📄 {rel_path} → {len(chunks)} DOCX chunks")
                    
    return documents


import requests
import time
from chromadb.api.types import Documents, EmbeddingFunction, Embeddings

# Removed custom RequestsHuggingFaceEmbeddingFunction to avoid ChromaDB signature conflicts

def build_knowledge_base():
    print("\n[RAG] Building MindBridge Knowledge Base (Multi-Format)...")
    documents = load_documents()
    print(f"\n[RAG] Total chunks loaded: {len(documents)}")

    if os.path.exists(CHROMA_DIR):
        try:
            shutil.rmtree(CHROMA_DIR)
            print("[RAG] Cleared previous ChromaDB directory")
        except Exception as e:
            print(f"[RAG] Directory reset note: {e}")

    hf_token = os.getenv("HF_TOKEN")
    if not hf_token:
        print("[WARNING] HF_TOKEN not found in environment. Skipping RAG build during Docker build.")
        return None

    try:
        print("[RAG] Initializing remote HuggingFaceEmbeddingFunction...")
        embedding_fn = chromadb.utils.embedding_functions.HuggingFaceEmbeddingFunction(
            api_key=hf_token,
            model_name="sentence-transformers/all-MiniLM-L6-v2"
        )
    except Exception as e:
        print(f"[WARNING] Could not initialize HuggingFaceEmbeddingFunction: {e}")
        print("[WARNING] Skipping RAG build.")
        return None

    from chromadb.config import Settings
    client = chromadb.PersistentClient(
        path=CHROMA_DIR,
        settings=Settings(anonymized_telemetry=False)
    )

    try:
        client.delete_collection(name=COLLECTION_NAME)
        print("[RAG] Cleared previous ChromaDB collection safely via API")
    except Exception:
        pass

    try:
        collection = client.create_collection(
            name=COLLECTION_NAME,
            embedding_function=embedding_fn,
            metadata={"hnsw:space": "cosine"},
        )
    except Exception as e:
        if "conflict" in str(e).lower() or "huggingface vs" in str(e).lower() or "already exists" in str(e).lower():
            print("[WARNING] Detected embedding function conflict. Wiping database to rebuild...")
            # If the API delete fails because of schema mismatch, we must wipe the folder
            try:
                if os.path.exists(CHROMA_DIR):
                    shutil.rmtree(CHROMA_DIR)
                client = chromadb.PersistentClient(
                    path=CHROMA_DIR,
                    settings=Settings(anonymized_telemetry=False)
                )
                collection = client.create_collection(
                    name=COLLECTION_NAME,
                    embedding_function=embedding_fn,
                    metadata={"hnsw:space": "cosine"},
                )
            except Exception as inner_e:
                print(f"[ERROR] Fatal rebuild error: {inner_e}")
                return None
        else:
            print(f"[ERROR] Failed to create collection: {e}")
            return None

    batch_size = 50
    for i in range(0, len(documents), batch_size):
        batch = documents[i:i + batch_size]
        
        # Add retry logic for Hugging Face API
        for attempt in range(5):
            try:
                collection.add(
                    ids=[d["id"] for d in batch],
                    documents=[d["text"] for d in batch],
                    metadatas=[{"source": d["source"], "concept": d["concept"], "technique": d["technique"]} for d in batch],
                )
                print(f"[RAG] Processed batch {i//batch_size + 1}")
                break
            except Exception as e:
                err_str = str(e).lower()
                if "503" in err_str or "502" in err_str or "timeout" in err_str or "loading" in err_str or "getaddrinfo" in err_str or "connecterror" in err_str:
                    print(f"[RAG] Model loading or network error ({e}). Retrying in 10s (attempt {attempt+1}/5)...")
                    import time
                    time.sleep(10)
                else:
                    raise e
        else:
            print("[ERROR] Failed to add batch after 5 attempts.")
            return None

    print(f"[RAG] Knowledge base built: {collection.count()} chunks stored")
    
    # Save timestamp marker file for smart deployment builds
    try:
        marker_path = os.path.join(CHROMA_DIR, ".built_at")
        with open(marker_path, "w") as f:
            f.write(str(os.path.getmtime(CHROMA_DIR)))
    except Exception:
        pass

    return collection


def is_knowledge_base_up_to_date() -> bool:
    """
    Check if the ChromaDB vector database exists and is newer than all document source files.
    """
    if not os.path.exists(CHROMA_DIR):
        return False

    marker_path = os.path.join(CHROMA_DIR, ".built_at")
    if not os.path.exists(marker_path):
        return False

    built_time = os.path.getmtime(marker_path)

    # Check docs directory timestamp
    if not os.path.exists(DOCS_DIR):
        return True

    for root, _, files in os.walk(DOCS_DIR):
        for file in files:
            fp = os.path.join(root, file)
            if os.path.getmtime(fp) > built_time:
                return False

    return True


_INIT_LOCK_FILE = os.path.join(os.path.dirname(__file__), "chroma_db", ".init_lock")
_init_completed = False


def ensure_knowledge_base_built(force: bool = False):
    """
    Ensures knowledge base is initialized on startup.
    - Logs clearly whether KB exists, was built, or failed.
    - Uses a file-based lock to prevent parallel worker initialization.
    - Validates the collection is readable after build.
    - Skips rebuild if database exists and is up-to-date.
    - Executes exactly once per process (guarded by module-level flag).
    """
    global _init_completed
    if _init_completed and not force:
        print("[RAG] Knowledge base already initialized in this process. Skipping.")
        return None

    # --- Check up-to-date first ---
    if not force and is_knowledge_base_up_to_date():
        print("[RAG] ✓ Knowledge Base is up to date. Skipping rebuild.")
        _init_completed = True
        return None

    # --- File-based lock to prevent simultaneous worker initialization ---
    lock_dir = os.path.dirname(_INIT_LOCK_FILE)
    if lock_dir:
        os.makedirs(lock_dir, exist_ok=True)
    if os.path.exists(_INIT_LOCK_FILE):
        print("[RAG] Another worker is already building the knowledge base. Waiting...")
        import time as _time
        for _ in range(60):  # Wait up to 60 seconds
            _time.sleep(1)
            if not os.path.exists(_INIT_LOCK_FILE):
                break
        print("[RAG] ✓ Knowledge Base initialized by another worker.")
        _init_completed = True
        return None

    # --- Acquire lock ---
    try:
        with open(_INIT_LOCK_FILE, "w") as lf:
            lf.write(str(os.getpid()))
    except Exception:
        pass  # Non-fatal — proceed without lock

    print("[RAG] ⚙ Knowledge Base missing or outdated. Building now...")
    collection = None
    try:
        collection = build_knowledge_base()
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print(f"[RAG] ✗ Knowledge Base build failed: {e}")
        print(f"[RAG] ✗ Diagnostic Traceback:\n{error_trace}")
    finally:
        # Always release lock
        try:
            os.remove(_INIT_LOCK_FILE)
        except Exception:
            pass

    if collection is not None:
        try:
            count = collection.count()
            print(f"[RAG] ✓ Knowledge Base ready — {count} chunks indexed.")
        except Exception as e:
            print(f"[RAG] ✗ Knowledge Base validation failed: {e}")
    else:
        hf_token = os.getenv("HF_TOKEN")
        if not hf_token:
            print("[RAG] ✗ Knowledge Base NOT built — HF_TOKEN is missing. RAG will be disabled until token is provided.")
        else:
            print("[RAG] ✗ Knowledge Base build returned no collection. Check logs above for errors.")

    _init_completed = True
    return collection


if __name__ == "__main__":
    ensure_knowledge_base_built()
    print("\n[RAG] Done! Multi-format knowledge base is ready.")

