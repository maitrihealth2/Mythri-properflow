"""
00_parse_documents.py
======================
Parses 'Knowledge 1 - Therapy .pdf' and 'Psychodynamic theory.docx'
into structured semantic chunks tagged by therapeutic domain.

Output: finetuning/data/raw_chunks.jsonl
Each line:
  {"source": "therapy_pdf", "domain": "CBT", "chunk": "...", "keywords": [...]}

Run:
  python finetuning/00_parse_documents.py
"""

import json
import re
import pathlib
import sys

BACKEND_ROOT = pathlib.Path(__file__).resolve().parent.parent
DOCS_DIR = BACKEND_ROOT.parent  # mythri-v5 root where PDFs live
OUTPUT_DIR = pathlib.Path(__file__).resolve().parent / "data"
OUTPUT_DIR.mkdir(exist_ok=True)

PDF_PATH = DOCS_DIR / "Knowledge 1 - Therapy .pdf"
DOCX_PATH = DOCS_DIR / "Psychodynamic theory.docx"

# Domain classification keywords
DOMAIN_KEYWORDS = {
    "CBT": [
        "cognitive", "behavioral", "thought", "automatic", "distortion",
        "restructuring", "belief", "schema", "cognitive triad", "beck",
        "behavioral activation", "thought record", "exposure", "socratic"
    ],
    "DBT": [
        "dialectical", "linehan", "distress tolerance", "mindfulness",
        "emotion regulation", "interpersonal effectiveness", "radical acceptance",
        "tipp", "dear man", "wise mind", "skills", "dbt"
    ],
    "ACT": [
        "acceptance", "commitment", "defusion", "values", "hexaflex",
        "experiential avoidance", "psychological flexibility", "present moment",
        "self-as-context", "observing self", "committed action", "hayes"
    ],
    "Psychodynamic": [
        "unconscious", "transference", "countertransference", "defense mechanism",
        "attachment", "object relations", "ego", "id", "superego", "freud",
        "psychoanalytic", "early childhood", "working through", "interpretation",
        "bowlby", "winnicott", "klein", "relational"
    ],
    "Humanistic": [
        "rogers", "person-centered", "unconditional positive regard", "empathy",
        "congruence", "self-actualization", "maslow", "gestalt", "existential",
        "humanistic", "client-centered"
    ],
    "Mindfulness": [
        "mindfulness", "meditation", "present moment", "awareness", "breathing",
        "body scan", "grounding", "mbsr", "mbct", "jon kabat-zinn", "vipassana"
    ],
    "Trauma": [
        "trauma", "ptsd", "emdr", "somatic", "dissociation", "flashback",
        "hyperarousal", "window of tolerance", "complex trauma", "van der kolk",
        "bessel", "nervous system"
    ],
    "Emotion": [
        "emotion", "affect", "feeling", "regulation", "dysregulation",
        "grief", "loss", "anger", "anxiety", "depression", "shame", "guilt"
    ],
}


def classify_domain(text: str) -> str:
    """Classify a text chunk into the most relevant therapeutic domain."""
    text_lower = text.lower()
    scores = {}
    for domain, keywords in DOMAIN_KEYWORDS.items():
        score = sum(1 for kw in keywords if kw.lower() in text_lower)
        scores[domain] = score
    best = max(scores, key=scores.get)
    return best if scores[best] > 0 else "General"


def extract_keywords(text: str) -> list[str]:
    """Extract relevant therapy keywords from a chunk."""
    found = []
    text_lower = text.lower()
    for domain, keywords in DOMAIN_KEYWORDS.items():
        for kw in keywords:
            if kw.lower() in text_lower and kw not in found:
                found.append(kw)
    return found[:10]


def parse_pdf(pdf_path: pathlib.Path) -> list[dict]:
    """Extract and chunk text from PDF."""
    try:
        import pypdf
    except ImportError:
        print("  [ERROR] pypdf not installed. Run: pip install pypdf")
        return []

    if not pdf_path.exists():
        print(f"  [WARNING] PDF not found at: {pdf_path}")
        return []

    print(f"  Reading: {pdf_path.name}")
    reader = pypdf.PdfReader(str(pdf_path))
    print(f"  Pages: {len(reader.pages)}")

    # Extract all text
    full_text = ""
    for page in reader.pages:
        page_text = page.extract_text() or ""
        full_text += page_text + "\n\n"

    return chunk_text(full_text, source="therapy_pdf")


def parse_docx(docx_path: pathlib.Path) -> list[dict]:
    """Extract and chunk text from DOCX."""
    try:
        import docx
    except ImportError:
        print("  [ERROR] python-docx not installed. Run: pip install python-docx")
        return []

    if not docx_path.exists():
        print(f"  [WARNING] DOCX not found at: {docx_path}")
        return []

    print(f"  Reading: {docx_path.name}")
    doc = docx.Document(str(docx_path))

    # Extract by paragraphs (preserving structure)
    full_text = "\n\n".join(para.text.strip() for para in doc.paragraphs if para.text.strip())
    print(f"  Paragraphs extracted: {len(doc.paragraphs)}")

    return chunk_text(full_text, source="psychodynamic_docx")


def chunk_text(text: str, source: str, min_words: int = 60, max_words: int = 300) -> list[dict]:
    """Split text into meaningful semantic chunks."""
    # Clean up whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]+', ' ', text)

    # Split on double newlines (paragraph boundaries)
    raw_paragraphs = [p.strip() for p in re.split(r'\n\n+', text) if p.strip()]

    chunks = []
    current_chunk = []
    current_word_count = 0

    for para in raw_paragraphs:
        word_count = len(para.split())
        if word_count < 10:
            continue  # Skip very short lines (headers, page numbers)

        if current_word_count + word_count > max_words and current_chunk:
            # Save current chunk
            chunk_text_str = " ".join(current_chunk)
            if len(chunk_text_str.split()) >= min_words:
                domain = classify_domain(chunk_text_str)
                keywords = extract_keywords(chunk_text_str)
                chunks.append({
                    "source": source,
                    "domain": domain,
                    "chunk": chunk_text_str,
                    "keywords": keywords,
                    "word_count": len(chunk_text_str.split()),
                })
            current_chunk = [para]
            current_word_count = word_count
        else:
            current_chunk.append(para)
            current_word_count += word_count

    # Don't forget the last chunk
    if current_chunk:
        chunk_text_str = " ".join(current_chunk)
        if len(chunk_text_str.split()) >= min_words:
            domain = classify_domain(chunk_text_str)
            keywords = extract_keywords(chunk_text_str)
            chunks.append({
                "source": source,
                "domain": domain,
                "chunk": chunk_text_str,
                "keywords": keywords,
                "word_count": len(chunk_text_str.split()),
            })

    print(f"  Chunks created: {len(chunks)}")
    return chunks


def main():
    print("=" * 60)
    print("Mythri Document Parser")
    print("=" * 60)

    all_chunks = []

    # 1. Parse PDF
    print("\n[1] Parsing therapy PDF...")
    pdf_chunks = parse_pdf(PDF_PATH)
    all_chunks.extend(pdf_chunks)

    # 2. Parse DOCX
    print("\n[2] Parsing psychodynamic theory DOCX...")
    docx_chunks = parse_docx(DOCX_PATH)
    all_chunks.extend(docx_chunks)

    if not all_chunks:
        print("\n[ERROR] No chunks were extracted. Check that documents exist at:")
        print(f"  {PDF_PATH}")
        print(f"  {DOCX_PATH}")
        sys.exit(1)

    # 3. Save
    output_path = OUTPUT_DIR / "raw_chunks.jsonl"
    with open(output_path, "w", encoding="utf-8") as f:
        for chunk in all_chunks:
            f.write(json.dumps(chunk, ensure_ascii=False) + "\n")

    # Summary
    print(f"\n[DONE] Extracted {len(all_chunks)} chunks:")
    domain_counts = {}
    for c in all_chunks:
        domain_counts[c["domain"]] = domain_counts.get(c["domain"], 0) + 1
    for domain, count in sorted(domain_counts.items(), key=lambda x: -x[1]):
        print(f"  {domain:20s}: {count} chunks")
    print(f"\nSaved to: {output_path}")
    print("Next: python finetuning/00b_extract_behaviors.py")


if __name__ == "__main__":
    main()
